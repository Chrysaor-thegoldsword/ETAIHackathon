"""Utility script to generate the architecture document and impact model.

Run this script with python to produce the docx reports used in the ET
Money Mentor submission.  The script uses matplotlib to draw a
simple architecture diagram and python‑docx to assemble a 1–2 page
document describing the agent roles and interactions, along with a
quantified impact model.
"""

import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


def create_architecture_diagram(output_path: str) -> None:
    """Create a block diagram of the multi‑agent architecture."""
    fig, ax = plt.subplots(figsize=(6, 4))
    ax.axis('off')

    # Define positions for each agent box
    positions = {
        "Conversation Agent": (0.1, 0.7),
        "Doc Extraction Agent": (0.1, 0.5),
        "Profile Builder": (0.1, 0.3),
        "Health Score Engine": (0.5, 0.7),
        "Goal Planner": (0.5, 0.55),
        "Tax Optimizer": (0.5, 0.4),
        "Portfolio Recommender": (0.5, 0.25),
        "Report Generator": (0.8, 0.5),
    }
    # Draw rectangles
    for name, (x, y) in positions.items():
        ax.add_patch(plt.Rectangle((x - 0.08, y - 0.05), 0.16, 0.08,
                                   edgecolor='black', facecolor='#f7f7f7'))
        ax.text(x, y, name, ha='center', va='center', fontsize=8, wrap=True)

    # Draw arrows
    def arrow(start, end):
        ax.annotate('', xy=end, xytext=start,
                    arrowprops=dict(arrowstyle='->', lw=1))

    # Conversation flows into doc extraction and profile builder
    arrow((0.18, 0.65), (0.18, 0.55))
    arrow((0.18, 0.45), (0.18, 0.35))
    # Profile builder outputs to analysis agents
    arrow((0.24, 0.3), (0.42, 0.7))  # to health
    arrow((0.24, 0.3), (0.42, 0.55))  # to goals
    arrow((0.24, 0.3), (0.42, 0.4))  # to tax
    arrow((0.24, 0.3), (0.42, 0.25))  # to portfolio
    # Analysis agents feed report generator
    arrow((0.66, 0.7), (0.74, 0.53))
    arrow((0.66, 0.55), (0.74, 0.53))
    arrow((0.66, 0.4), (0.74, 0.53))
    arrow((0.66, 0.25), (0.74, 0.53))
    # Conversation agent also gets final report
    arrow((0.78, 0.47), (0.2, 0.73))

    fig.tight_layout()
    fig.savefig(output_path)
    plt.close(fig)


def build_architecture_doc(diagram_path: str, output_docx: str) -> None:
    """Create a DOCX file describing the architecture and impact model."""
    doc = Document()
    doc.add_heading('ET Money Mentor Architecture', level=1)
    doc.add_paragraph(
        'ET Money Mentor uses a multi‑agent system to guide users through a complete financial planning workflow. '
        'Each agent is responsible for a specific domain expertise, and the orchestrator (Conversation Agent) ensures '
        'a smooth, human‑friendly dialogue. The system is designed for extensibility and fault tolerance.')

    doc.add_heading('Agent Roles', level=2)
    doc.add_paragraph('Conversation Agent – orchestrates the dialogue, collects user inputs, triggers downstream agents and delivers a final summary.')
    doc.add_paragraph('Document Extraction Agent – parses uploaded documents (ITR/Form 16/MF statements) to pre‑fill financial fields.')
    doc.add_paragraph('Profile Builder – normalizes all collected data into a structured profile for analysis.')
    doc.add_paragraph('Health Score Engine – computes a six‑factor wellbeing score and highlights red flags.')
    doc.add_paragraph('Goal Planner – translates user goals into monthly savings targets and milestones.')
    doc.add_paragraph('Tax Optimizer – compares the old vs new regimes and recommends the most efficient path.')
    doc.add_paragraph('Portfolio Recommender – proposes an asset allocation and sample instruments based on risk appetite.')
    doc.add_paragraph('Report Generator – compiles all outputs into a single report for the user.')

    doc.add_heading('Architecture Diagram', level=2)
    doc.add_picture(diagram_path, width=Inches(5.5))
    doc.add_paragraph('Figure 1 – Multi‑agent architecture. Arrows indicate the flow of information between agents.')

    doc.add_heading('Error Handling & Guardrails', level=2)
    doc.add_paragraph(
        'The orchestrator validates user inputs and ensures that the downstream agents receive the right data. '
        'If an agent fails (e.g., document parsing returns an error), the conversation agent requests clarification or offers '
        'an alternative path. All decisions are auditable, and the system refrains from performing any high‑risk financial transactions.')

    # Impact model section
    doc.add_heading('Impact Model', level=1)
    doc.add_paragraph(
        'Assuming a typical salaried user with annual income of ₹6 Lakh and three financial goals (emergency fund, home down payment, retirement), '
        'ET Money Mentor saves users both time and money.')
    doc.add_paragraph(
        '- **Time saved**: Instead of spending ~5 hours researching tax rules, investment options and budget calculators, '
        'the user completes the onboarding and receives a plan within 10 minutes.  This is a **95% reduction** in planning time.')
    doc.add_paragraph(
        '- **Cost reduced**: Financial planners typically charge ₹25,000 per year for basic advisory【47679722380880†L330-L352】.  Assuming 50% of users adopt the DIY plan and avoid paying a planner, '
        'the system saves ₹12,500 per user annually.')
    doc.add_paragraph(
        '- **Revenue opportunity for ET**: Out of 100,000 monthly ET users exploring investment articles, if 8% start the mentor, 50% finish and 5% purchase a recommended product, the monetizable funnel yields 200 paying users. '
        'If ET earns ₹1,000 commission per conversion, this translates to **₹2 Lakh in monthly revenue** (₹24 Lakh annually).')
    doc.add_paragraph('These estimates are illustrative; you should tailor the numbers to your audience and include your own assumptions.')

    doc.add_paragraph('This document and diagram were generated programmatically as part of the hackathon submission.')
    doc.save(output_docx)


if __name__ == '__main__':
    # Paths relative to the project root
    diag_path = os.path.join('et_money_mentor', 'architecture_diagram.png')
    doc_path = os.path.join('et_money_mentor', 'architecture.docx')
    create_architecture_diagram(diag_path)
    build_architecture_doc(diag_path, doc_path)
    print(f"Generated {diag_path} and {doc_path}")
